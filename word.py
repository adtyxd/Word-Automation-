import os
import subprocess
import docx
import requests
import re
import platform


# Set your API key here

API_KEY = "your API key"

# URL for your API
API_URL = "https://api.mistral.ai/v1/chat/completions"

def get_mistral_response(prompt):
    headers = {
        "Authorization": f"Bearer {API_KEY}",
        "Content-Type": "application/json"
    }

    data = {
        "model": "mistral-large-2407",  # Change model 
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.7
    }

    response = requests.post(API_URL, json=data, headers=headers)
    
    if response.status_code == 200:
        return response.json()["choices"][0]["message"]["content"]
    else:
        return f"Error: {response.status_code}, {response.text}"

def sanitize_filename(title):
    """Removes invalid characters and AI artifacts from filenames."""
    title = re.sub(r'[<>:"/\\|?*]', '', title).strip()
    title = re.sub(r'\s*end_of_turn\s*', '', title, flags=re.IGNORECASE).strip()  
    return title

def create_word_file():
    """Creates a Word document with AI-generated content and a clean, centered title."""
    
    query = input("Enter command (type 'word' to start): ").strip().lower()
    
    if "make a word file" in query or "create word file" in query or "word" in query:
        content_description = input("What should be in the document? ")

        
        content = get_mistral_response(content_description)
        
        if content:
            
            title_prompt = f"Provide a concise, professional title (max 4 words) for the following PROMPT:\n\n{content_description}"
            title = get_mistral_response(title_prompt)

            if not title:
                title = "Untitled Document"

            
            safe_title = sanitize_filename(title)

           
            doc = docx.Document()

           
           
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(safe_title)
            title_run.bold = True  
            title_run.font.size = docx.shared.Pt(16)  
            title_para.alignment = 1  

            doc.add_paragraph("\n") 

            
            doc.add_paragraph(content.strip())

            
            file_path = os.path.abspath(f"{safe_title}.docx")
            doc.save(file_path)
            print(f"Word file '{file_path}' created successfully.")

            
            if os.path.exists(file_path):
                print(f"Opening file: {file_path}")
                if platform.system() == "Darwin":
                    subprocess.run(["open", file_path], check=True)
                elif platform.system() == "Windows":
                    os.startfile(file_path)
                elif platform.system() == "Linux":
                    subprocess.run(["xdg-open", file_path], check=True)
            else:
                print("Error: File was not created!")

        else:
            print("AI failed to n content. Please try again.")

create_word_file()
